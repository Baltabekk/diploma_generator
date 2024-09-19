import os
import google.generativeai as genai
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import DiplomaRequest, Feedback
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

GENAI_API_KEY = 'AIzaSyCtrFiYRihVUm_L58vS-c_8MEyZX7VLLv0'
genai.configure(api_key=GENAI_API_KEY)
model = genai.GenerativeModel('gemini-pro')

def home(request):
    return render(request, 'home.html')

def generate(request):
    if request.method == 'POST':
        topic = request.POST.get('topic')
        size = int(request.POST.get('size'))
        
        diploma_request = DiplomaRequest.objects.create(topic=topic, size=size)
        
        content = generate_content(topic, size)
        document = create_document(topic, content)
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={topic.replace(" ", "_")}.docx'
        document.save(response)
        
        return response
    
    return render(request, 'generate.html')

def generate_content(topic, size):
    if size == 40:
        prompt = f"Create a detailed outline for a thesis on '{topic}'. Include an introduction, 4 main sections with 4 subsections each, and a conclusion."
    elif size == 60:
        prompt = f"Create a detailed outline for a thesis on '{topic}'. Include an introduction, 6 main sections with 4 subsections each, and a conclusion."
    else:  # 100
        prompt = f"Create a detailed outline for a thesis on '{topic}'. Include an introduction, 7 main sections with 5 subsections each, and a conclusion."
    
    response = model.generate_content(prompt)
    outline = response.text
    
    sections = ["Introduction"] + [section.strip() for section in outline.split('\n') if section.strip()] + ["Conclusion"]
    content = []
    
    for section in sections:
        section_prompt = f"Write a detailed text (about 500 words) for the section '{section}' of a thesis on '{topic}'."
        section_response = model.generate_content(section_prompt)
        content.append((section, section_response.text))
    
    return content

def create_document(topic, content):
    doc = Document()
    styles = doc.styles
    toc_style = styles.add_style('TOC', WD_STYLE_TYPE.PARAGRAPH)
    toc_style.font.size = Pt(12)
    toc_style.font.name = 'Times New Roman'
    
    doc.add_heading(f"Thesis on: {topic}", level=0)
    doc.add_paragraph("Table of Contents")
    
    for section, _ in content:
        doc.add_paragraph(section, style='TOC')
    
    doc.add_page_break()
    
    for section, text in content:
        doc.add_heading(section, level=1)
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            doc.add_paragraph(para)
    
    return doc

def feedback(request):
    if request.method == 'POST':
        content = request.POST.get('content')
        Feedback.objects.create(content=content)
        return redirect('home')
    
    feedbacks = Feedback.objects.all().order_by('-created_at')[:5]
    return render(request, 'feedback.html', {'feedbacks': feedbacks})

